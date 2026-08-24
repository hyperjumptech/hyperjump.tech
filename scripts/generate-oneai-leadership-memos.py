#!/usr/bin/env python3
"""Generate bilingual one-page OneAI leadership-memo DOCX files."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

NAVY = RGBColor(0x0A, 0x25, 0x40)
GRAY = RGBColor(0x42, 0x54, 0x66)
MUTED = RGBColor(0x88, 0x98, 0xAA)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
RULE = "0A2540"
LIGHT_FILL = "EEF5FB"
AMBER_FILL = "FFF6D8"
REPO_ROOT = Path(__file__).resolve().parent.parent
OUTPUT_DIR = REPO_ROOT / "public" / "documents"

EN = {
    "output": "oneai-leadership-memo-en.docx",
    "core_title": "OneAI leadership memo",
    "core_subject": "Recommendation to evaluate OneAI",
    "kicker": "HYPERJUMP TECHNOLOGY",
    "brand": "OneAI",
    "doc_type": "Internal memo  ·  one page  ·  send to your manager or leadership",
    "instruction": (
        "How to use this file: replace every [bracketed] field, attach the "
        "OneAI two-page PDF overview, then delete this yellow box before you send."
    ),
    "fields": [
        ("To", "[Name], [Title]"),
        ("From", "[Your name], [Role / department]"),
        ("Date", "[Date]"),
        (
            "Subject",
            "Recommendation to evaluate OneAI as our company AI platform",
        ),
        ("Attachments", "OneAI two-page overview (PDF)"),
    ],
    "request_heading": "1. Request",
    "request": (
        "I recommend we evaluate OneAI from Hyperjump Technology as our "
        "company-wide AI platform, or at minimum schedule a 30-minute briefing. "
        "This is not about replacing personal ChatGPT for a few people. It is "
        "about governing AI now that 20-40 of us already use it."
    ),
    "why_heading": "2. Why this matters now",
    "why_intro": (
        "Teams are already using AI. Personal subscriptions and informal "
        "stipends do not give us:"
    ),
    "why_points": [
        (
            "Control",
            "one dashboard, usage reports, and a way to choose which models are available.",
        ),
        (
            "A single budget",
            "one invoice instead of 20-40 subscriptions we cannot verify.",
        ),
        (
            "Data that stays with us",
            "installed on our premises (Hyperjump maintains it) or an isolated instance for us only. Zero Data Retention: data is not kept for model training.",
        ),
        (
            "Real capability",
            "office files, internal knowledge, code execution, and more than one model vendor.",
        ),
    ],
    "cost": (
        "At 40 users, OneAI is Rp12.400.000 per month (USD 300 AI credit "
        "included; billed quarterly). ChatGPT Business is about Rp16 million "
        "monthly or Rp12.8 million annually. The prices are close. The case is "
        "one platform across vendors, not cheaper chat. The attached PDF has "
        "the full comparison. Tax is extra."
    ),
    "caveat": (
        "If we have fewer than 20 people and only need chat, ChatGPT Plus or "
        "Business is usually enough and cheaper."
    ),
    "next_heading": "3. Proposed next step",
    "next": (
        "Please approve a 30-minute call with Hyperjump "
        "(solution@hyperjump.tech) to confirm fit, onboarding, and commercial "
        "terms. I am happy to coordinate and report back."
    ),
    "signoff": "Respectfully,",
    "signature_name": "[Your name]",
    "signature_role": "[Role / department]",
    "signature_contact": "[Phone / email]",
    "footer": (
        "Vendor: Hyperjump Technology · PT Artha Rajamas Mandiri · Jakarta  ·  "
        "solution@hyperjump.tech  ·  https://hyperjump.tech/en/oneai"
    ),
}

ID = {
    "output": "oneai-leadership-memo-id.docx",
    "core_title": "Memo pimpinan OneAI",
    "core_subject": "Usulan untuk mengevaluasi OneAI",
    "kicker": "HYPERJUMP TECHNOLOGY",
    "brand": "OneAI",
    "doc_type": "Memo internal  ·  satu halaman  ·  kirim ke manajer atau pimpinan",
    "instruction": (
        "Cara memakai berkas ini: ganti setiap isian [dalam kurung siku], "
        "lampirkan ringkasan PDF OneAI dua halaman, lalu hapus kotak kuning ini sebelum dikirim."
    ),
    "fields": [
        ("Kepada", "[Nama Bapak/Ibu], [Jabatan]"),
        ("Dari", "[Nama Anda], [Jabatan / divisi]"),
        ("Tanggal", "[Tanggal]"),
        (
            "Perihal",
            "Usulan untuk mengevaluasi OneAI sebagai platform AI perusahaan",
        ),
        ("Lampiran", "Ringkasan OneAI dua halaman (PDF)"),
    ],
    "request_heading": "1. Permohonan",
    "request": (
        "Saya mengusulkan agar perusahaan mengevaluasi OneAI dari Hyperjump "
        "Technology sebagai platform AI terpadu, atau setidaknya menjadwalkan "
        "briefing 30 menit. Ini bukan usulan mengganti ChatGPT pribadi untuk "
        "beberapa orang. Ini usulan mengatur AI begitu 20-40 orang di "
        "perusahaan sudah memakainya."
    ),
    "why_heading": "2. Mengapa ini penting sekarang",
    "why_intro": (
        "Tim sudah memakai AI. Langganan pribadi dan <i>allowance</i> informal "
        "tidak memberi kita:"
    ),
    "why_points": [
        (
            "Kontrol",
            "satu dashboard, laporan pemakaian, dan cara memilih model yang tersedia.",
        ),
        (
            "Satu anggaran",
            "satu tagihan, bukan 20-40 langganan yang tidak terverifikasi.",
        ),
        (
            "Data tetap di sisi kita",
            "dipasang on-premise (Hyperjump merawatnya) atau instance terisolasi khusus kita. ZDR: data tidak disimpan untuk training model.",
        ),
        (
            "Kapabilitas yang nyata",
            "file kantor, pengetahuan internal, eksekusi kode, dan lebih dari satu vendor model.",
        ),
    ],
    "cost": (
        "Di 40 pengguna, OneAI Rp12.400.000 per bulan (kredit AI USD 300 "
        "termasuk; ditagih per kuartal). ChatGPT Business sekitar Rp16 juta "
        "bulanan atau Rp12,8 juta tahunan. Harganya berdekatan. Argumennya "
        "satu platform lintas vendor, bukan chat yang lebih murah. Tabel "
        "lengkap ada di PDF terlampir. Pajak belum termasuk."
    ),
    "caveat": (
        "Jika kita kurang dari 20 orang dan kebutuhannya hanya chat, ChatGPT "
        "Plus atau Business biasanya cukup dan lebih murah."
    ),
    "next_heading": "3. Usulan langkah berikutnya",
    "next": (
        "Mohon persetujuan panggilan 30 menit dengan Hyperjump "
        "(solution@hyperjump.tech) untuk konfirmasi kesesuaian, onboarding, "
        "dan ketentuan komersial. Saya siap mengkoordinasikan dan melapor kembali."
    ),
    "signoff": "Hormat saya,",
    "signature_name": "[Nama Anda]",
    "signature_role": "[Jabatan / divisi]",
    "signature_contact": "[Telepon / email]",
    "footer": (
        "Vendor: Hyperjump Technology · PT Artha Rajamas Mandiri · Jakarta  ·  "
        "solution@hyperjump.tech  ·  https://hyperjump.tech/id/oneai"
    ),
}


def set_run_font(run, *, name="Calibri", size=11, bold=False, color=NAVY, italic=False):
    """Apply a consistent font to a python-docx run."""
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_paragraph_spacing(paragraph, *, before=0, after=8, line=1.08):
    """Set paragraph spacing in points."""
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def shade_cell(cell, fill):
    """Set a table cell background hex fill."""
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:shd"))
    if existing is not None:
        tc_pr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_borders(cell, **edges):
    """Set individual cell borders. Each edge is (sz, color)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge, (sz, color) in edges.items():
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(sz))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        tc_borders.append(element)


def prevent_table_indent(table):
    """Remove default table left indent so full-width tables align with body text."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_cm):
    """Assign fixed column widths in centimeters."""
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for index, width in enumerate(widths_cm):
            row.cells[index].width = Cm(width)


ITALIC_TAG = re.compile(r"<i>(.*?)</i>")


def add_text_with_placeholders(paragraph, text, font):
    """Add a run, highlighting [placeholder] tokens in amber."""
    remaining = text
    while remaining:
        start = remaining.find("[")
        end = remaining.find("]", start + 1) if start != -1 else -1
        if start == -1 or end == -1:
            run = paragraph.add_run(remaining)
            set_run_font(run, **font)
            return
        if start > 0:
            run = paragraph.add_run(remaining[:start])
            set_run_font(run, **font)
        token = remaining[start : end + 1]
        run = paragraph.add_run(token)
        highlight = {**font, "bold": True, "color": RGBColor(0x8A, 0x5A, 0x00)}
        set_run_font(run, **highlight)
        rpr = run._element.get_or_add_rPr()
        hl = OxmlElement("w:highlight")
        hl.set(qn("w:val"), "yellow")
        rpr.append(hl)
        remaining = remaining[end + 1 :]


def add_text(paragraph, text, **font):
    """Add runs, italicizing <i> fragments and highlighting [placeholder] tokens."""
    pos = 0
    for match in ITALIC_TAG.finditer(text):
        if match.start() > pos:
            add_text_with_placeholders(paragraph, text[pos : match.start()], font)
        add_text_with_placeholders(
            paragraph, match.group(1), {**font, "italic": True}
        )
        pos = match.end()
    if pos < len(text):
        add_text_with_placeholders(paragraph, text[pos:], font)


def add_body_paragraph(doc, text, *, after=6, italic=False, size=11):
    """Add a justified body paragraph."""
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=after)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_text(paragraph, text, size=size, color=GRAY, italic=italic)
    return paragraph


def add_heading_paragraph(doc, text):
    """Add a section heading sized for a one-page memo."""
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=8, after=3)
    run = paragraph.add_run(text)
    set_run_font(run, size=12, bold=True, color=NAVY)
    return paragraph


def add_bullet(doc, title, text):
    """Add a bold-lead bullet."""
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(paragraph, after=2)
    title_run = paragraph.add_run(f"{title}: ")
    set_run_font(title_run, size=10.5, bold=True, color=NAVY)
    add_text(paragraph, text, size=10.5, color=GRAY)


def add_header_bar(doc, copy):
    """Add a compact OneAI brand header."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    prevent_table_indent(table)
    set_table_widths(table, [16.6])
    cell = table.cell(0, 0)
    shade_cell(cell, "0A2540")
    set_cell_borders(
        cell,
        top=("0", "0A2540"),
        left=("0", "0A2540"),
        bottom=("0", "0A2540"),
        right=("0", "0A2540"),
    )
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    kicker = cell.paragraphs[0]
    kicker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(kicker, before=2, after=0)
    kicker_run = kicker.add_run(copy["kicker"])
    set_run_font(kicker_run, size=8, bold=True, color=RGBColor(0x00, 0xD4, 0xAA))

    brand = cell.add_paragraph()
    set_paragraph_spacing(brand, before=0, after=0)
    brand_run = brand.add_run(copy["brand"])
    set_run_font(brand_run, size=18, bold=True, color=WHITE)

    doc_type = cell.add_paragraph()
    set_paragraph_spacing(doc_type, before=1, after=2)
    type_run = doc_type.add_run(copy["doc_type"])
    set_run_font(type_run, size=9, color=RGBColor(0xC8, 0xD9, 0xEB))

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=6)


def add_instruction_box(doc, copy):
    """Add the yellow how-to box the sender should delete."""
    table = doc.add_table(rows=1, cols=1)
    prevent_table_indent(table)
    set_table_widths(table, [16.6])
    cell = table.cell(0, 0)
    shade_cell(cell, AMBER_FILL)
    set_cell_borders(
        cell,
        top=("8", "E6C200"),
        left=("8", "E6C200"),
        bottom=("8", "E6C200"),
        right=("8", "E6C200"),
    )
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, before=1, after=1)
    add_text(paragraph, copy["instruction"], size=9, color=RGBColor(0x6B, 0x53, 0x00))
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=4)


def add_field_table(doc, copy):
    """Add To/From/Date/Subject fields."""
    table = doc.add_table(rows=len(copy["fields"]), cols=2)
    prevent_table_indent(table)
    set_table_widths(table, [3.4, 13.2])
    for index, (label, value) in enumerate(copy["fields"]):
        label_cell = table.cell(index, 0)
        value_cell = table.cell(index, 1)
        shade_cell(label_cell, LIGHT_FILL)
        for cell in (label_cell, value_cell):
            set_cell_borders(
                cell,
                top=("4", "C8D9EB"),
                left=("4", "C8D9EB"),
                bottom=("4", "C8D9EB"),
                right=("4", "C8D9EB"),
            )
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        label_p = label_cell.paragraphs[0]
        set_paragraph_spacing(label_p, before=1, after=1)
        label_run = label_p.add_run(label)
        set_run_font(label_run, size=9, bold=True, color=NAVY)
        value_p = value_cell.paragraphs[0]
        set_paragraph_spacing(value_p, before=1, after=1)
        add_text(value_p, value, size=10.5, color=NAVY)
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=2)


def add_signature(doc, copy):
    """Add the sender sign-off with placeholder fields."""
    signoff = doc.add_paragraph()
    set_paragraph_spacing(signoff, before=8, after=8)
    add_text(signoff, copy["signoff"], size=11, color=GRAY)
    for key, size, bold in (
        ("signature_name", 12, True),
        ("signature_role", 10.5, False),
        ("signature_contact", 10, False),
    ):
        paragraph = doc.add_paragraph()
        set_paragraph_spacing(paragraph, before=0, after=0)
        add_text(paragraph, copy[key], size=size, bold=bold, color=NAVY)


def add_footer_line(doc, copy):
    """Add the vendor contact footer."""
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=10, after=0)
    paragraph.paragraph_format.border_top = None
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "8")
    top.set(qn("w:space"), "6")
    top.set(qn("w:color"), RULE)
    p_bdr.append(top)
    p_pr.append(p_bdr)
    add_text(paragraph, copy["footer"], size=8, color=MUTED)


def build_document(copy):
    """Build one locale's one-page leadership memo and write it to public/documents."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    doc.core_properties.title = copy["core_title"]
    doc.core_properties.subject = copy["core_subject"]
    doc.core_properties.author = "Hyperjump Technology"
    doc.core_properties.category = "OneAI"

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = GRAY

    add_header_bar(doc, copy)
    add_instruction_box(doc, copy)
    add_field_table(doc, copy)
    add_heading_paragraph(doc, copy["request_heading"])
    add_body_paragraph(doc, copy["request"])
    add_heading_paragraph(doc, copy["why_heading"])
    add_body_paragraph(doc, copy["why_intro"], after=4)
    for title, text in copy["why_points"]:
        add_bullet(doc, title, text)
    add_body_paragraph(doc, copy["cost"], after=4)
    add_body_paragraph(doc, copy["caveat"], after=4, italic=True, size=10)
    add_heading_paragraph(doc, copy["next_heading"])
    add_body_paragraph(doc, copy["next"], after=4)
    add_signature(doc, copy)
    add_footer_line(doc, copy)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    path = OUTPUT_DIR / copy["output"]
    doc.save(path)
    return path


def main():
    """Generate English and Indonesian one-page leadership memos."""
    for copy in (EN, ID):
        path = build_document(copy)
        print(f"Wrote {path.relative_to(REPO_ROOT)}")


if __name__ == "__main__":
    main()
